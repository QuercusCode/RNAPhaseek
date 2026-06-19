"""Render docs/RNAPhaseek_manuscript.md into a clean, submission-ready .docx
(NAR-style: title, section headings, justified 11pt body, numbered references).
Run with the base conda python (has python-docx):
  /opt/homebrew/Caskroom/mambaforge/base/bin/python scripts/reporting/build_manuscript_docx.py
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "docs/RNAPhaseek_manuscript.md"
OUT = "docs/RNAPhaseek_manuscript.docx"

doc = Document()
# base style
st = doc.styles["Normal"].font; st.name = "Calibri"; st.size = Pt(11)
for h, sz in (("Heading 1", 15), ("Heading 2", 12.5)):
    f = doc.styles[h].font; f.size = Pt(sz); f.color.rgb = RGBColor(0, 0, 0); f.name = "Calibri"


def add_runs(p, text):
    """Inline **bold**, *italic*, `code`."""
    text = text.replace("\\*", "*").replace("\\_", "_")
    pat = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")
    pos = 0
    for m in pat.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        t = m.group(0)
        if t.startswith("**"):
            r = p.add_run(t[2:-2]); r.bold = True
        elif t.startswith("`"):
            r = p.add_run(t[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        else:
            r = p.add_run(t[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


lines = open(SRC).read().split("\n")
seen_section = False   # front matter (centered) until the first "## "
i = 0
while i < len(lines):
    ln = lines[i].rstrip()
    if not ln.strip():
        i += 1; continue
    if ln.startswith("## "):
        seen_section = True
        doc.add_heading(ln[3:].strip(), level=1)
    elif ln.startswith("### "):
        doc.add_heading(ln[4:].strip(), level=2)
    elif ln.startswith("# "):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(ln[2:].strip()); r.bold = True; r.font.size = Pt(17)
    elif ln.strip() == "---":
        pass  # horizontal rule -> skip
    elif re.match(r"^\d+\.\s", ln):                       # numbered list (references)
        p = doc.add_paragraph(style="List Number"); add_runs(p, re.sub(r"^\d+\.\s", "", ln))
    elif ln.startswith("- ") or ln.startswith("* "):       # bullet
        p = doc.add_paragraph(style="List Bullet"); add_runs(p, ln[2:])
    else:                                                   # body paragraph
        p = doc.add_paragraph()
        if not seen_section:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, ln)
    i += 1

doc.save(OUT)
print(f"wrote {OUT}  ({len(doc.paragraphs)} paragraphs)")
